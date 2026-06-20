"""Document parsers and chunkers for building RAG knowledge bases.

Supports PDF, Markdown, Word (.docx), and CSV documents with
configurable chunking strategies.
"""

from __future__ import annotations

import csv
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Data types
# ---------------------------------------------------------------------------


@dataclass
class ParsedChunk:
    """A single chunk after parsing and splitting a document."""

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    chunk_index: int = 0
    total_chunks: int = 0


@dataclass
class ParsedDocument:
    """Result of parsing a document into chunks."""

    source: str
    title: str = ""
    chunks: list[ParsedChunk] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n".join(c.text for c in self.chunks)


# ---------------------------------------------------------------------------
# Text chunker
# ---------------------------------------------------------------------------


class TextChunker:
    """Split long text into overlapping chunks by character / token boundaries.

    Strategies (``mode``):
    - ``"fixed"`` — split at every *chunk_size* characters, with *chunk_overlap*.
    - ``"recursive"`` — split on natural separators (``\\n\\n``, ``\\n``, ``.``,
      `` ``, ``\\s``) from coarsest to finest until chunks fit.
    """

    _RECURSIVE_SEPARATORS = ["\n\n", "\n", ". ", "。", " ", ""]

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        mode: str = "recursive",
    ) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.mode = mode

    def split(self, text: str) -> list[str]:
        if self.mode == "fixed":
            return self._split_fixed(text)
        return self._split_recursive(text)

    # -- fixed ----------------------------------------------------------------

    def _split_fixed(self, text: str) -> list[str]:
        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
            if start >= len(text):
                break
        return chunks

    # -- recursive ------------------------------------------------------------

    def _split_recursive(self, text: str) -> list[str]:
        return self._recursive_split(text, self._RECURSIVE_SEPARATORS)

    def _recursive_split(self, text: str, separators: list[str]) -> list[str]:
        """Keep splitting until every piece fits in chunk_size."""
        if len(text) <= self.chunk_size:
            return [text] if text else []

        sep = separators[0] if separators else ""
        if not sep:
            # Final fallback: hard character split
            return self._split_fixed(text)

        parts = text.split(sep)
        merged: list[str] = []
        current: list[str] = []

        for part in parts:
            trial = sep.join(current + [part]) if current else part
            if len(trial) <= self.chunk_size:
                current.append(part)
            else:
                # Flush current batch
                if current:
                    merged.extend(self._merge_with_overlap(current, sep))
                # Handle oversized part with next separator
                sub = self._recursive_split(part, separators[1:])
                if sub:
                    # Merge last from previous and first from sub with overlap
                    if merged and sub:
                        last = merged[-1]
                        first = sub[0]
                        if len(last) < self.chunk_overlap:
                            overlap_text = last[-self.chunk_overlap :]
                            sub[0] = overlap_text + first
                    merged.extend(sub)
                current = []
        if current:
            merged.extend(self._merge_with_overlap(current, sep))

        return [m for m in merged if m]

    def _merge_with_overlap(self, parts: list[str], sep: str) -> list[str]:
        """Merge parts into chunks with overlap between chunks."""
        chunks: list[str] = []
        current_parts: list[str] = []
        current_len = 0

        for part in parts:
            sep_len = len(sep) if current_parts else 0
            if current_len + sep_len + len(part) <= self.chunk_size:
                current_parts.append(part)
                current_len += sep_len + len(part)
            else:
                if current_parts:
                    chunk = sep.join(current_parts)
                    chunks.append(chunk)
                # Start new chunk, carrying overlap from end of previous
                if chunks and self.chunk_overlap > 0:
                    prev = chunks[-1]
                    if len(prev) > self.chunk_overlap:
                        overlap = prev[-self.chunk_overlap :]
                        current_parts = [overlap, part]
                        current_len = len(overlap) + len(sep) + len(part)
                    else:
                        current_parts = [part]
                        current_len = len(part)
                else:
                    current_parts = [part]
                    current_len = len(part)

        if current_parts:
            chunks.append(sep.join(current_parts))

        return chunks


# ---------------------------------------------------------------------------
# Abstract document parser
# ---------------------------------------------------------------------------


class DocumentParser(ABC):
    """Base class for document format parsers."""

    @abstractmethod
    def parse(self, path: str | Path) -> str: ...

    def name(self) -> str:
        return self.__class__.__name__


# ---------------------------------------------------------------------------
# Format-specific parsers
# ---------------------------------------------------------------------------


class MarkdownParser(DocumentParser):
    """Parse Markdown files, stripping YAML front-matter."""

    def parse(self, path: str | Path) -> str:
        text = Path(path).read_text(encoding="utf-8")
        # Strip YAML front-matter
        text = re.sub(r"^---\n.*?\n---\n", "", text, flags=re.DOTALL)
        return text.strip()


class CSVParser(DocumentParser):
    """Parse CSV files into a readable text representation."""

    def __init__(self, delimiter: str = ",", max_rows: int = 0) -> None:
        self.delimiter = delimiter
        self.max_rows = max_rows

    def parse(self, path: str | Path) -> str:
        with Path(path).open(encoding="utf-8", newline="") as f:
            reader = csv.reader(f, delimiter=self.delimiter)
            rows = list(reader)
        if self.max_rows and len(rows) > self.max_rows:
            rows = rows[: self.max_rows] + [["... truncated"]]
        if not rows:
            return ""
        # Format as pipe-style table
        header = rows[0]
        lines = [" | ".join(header), " | ".join("---" for _ in header)]
        for row in rows[1:]:
            padded = row + [""] * (len(header) - len(row))
            lines.append(" | ".join(padded[: len(header)]))
        return "\n".join(lines)


class PDFParser(DocumentParser):
    """Parse PDF files. Requires ``pypdf`` (``pip install morainet-ai[pdf]``)."""

    def parse(self, path: str | Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError(
                "PDFParser requires pypdf. Install with: pip install morainet-ai[pdf]"
            ) from exc

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)


class WordParser(DocumentParser):
    """Parse Word (.docx) files. Requires ``python-docx`` (``pip install morainet-ai[docx]``)."""

    def parse(self, path: str | Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "WordParser requires python-docx. Install with: pip install morainet-ai[docx]"
            ) from exc

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Unified document loader
# ---------------------------------------------------------------------------


_PARSER_REGISTRY: dict[str, type[DocumentParser]] = {
    ".md": MarkdownParser,
    ".markdown": MarkdownParser,
    ".txt": MarkdownParser,
    ".csv": CSVParser,
    ".tsv": CSVParser,
}


def _get_parser(extension: str) -> DocumentParser:
    ext = extension.lower()
    cls = _PARSER_REGISTRY.get(ext)
    if cls is not None:
        return cls()
    raise ValueError(
        f"No parser registered for '{ext}'. Supported formats: {', '.join(sorted(_PARSER_REGISTRY))}. "
        "PDF and Word require optional extras: morainet-ai[pdf] / morainet-ai[docx]"
    )


def register_parser(extension: str, parser_cls: type[DocumentParser]) -> None:
    """Register a custom parser for a file extension."""
    ext = extension.lower()
    if not ext.startswith("."):
        ext = f".{ext}"
    _PARSER_REGISTRY[ext] = parser_cls


class DocumentLoader:
    """Load documents from files or directories, parse, chunk, and embed.

    Example::

        loader = DocumentLoader(chunk_size=800, embedder=OpenAIEmbedder())
        docs = loader.load_directory("docs/", glob="**/*.md")

        # Build a knowledge base in one go
        kb = await loader.build_knowledge_base("docs/", store=ChromaStore(path="./kb"))
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        chunk_mode: str = "recursive",
        embedder: Any | None = None,
    ) -> None:
        self.chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap, mode=chunk_mode)
        self.embedder = embedder

    # -- loading ---------------------------------------------------------------

    def load_file(self, path: str | Path) -> ParsedDocument:
        """Parse and chunk a single file."""
        path = Path(path)
        ext = path.suffix
        if ext in (".pdf", ".docx"):
            # Late-register optional parsers
            try:
                if ext == ".pdf":
                    _PARSER_REGISTRY.setdefault(".pdf", PDFParser)
                elif ext == ".docx":
                    _PARSER_REGISTRY.setdefault(".docx", WordParser)
            except Exception:
                pass

        parser = _get_parser(ext)
        raw_text = parser.parse(path)
        chunks_text = self.chunker.split(raw_text)
        total = len(chunks_text)
        chunks = [
            ParsedChunk(text=t, chunk_index=i, total_chunks=total, metadata={"source": str(path)})
            for i, t in enumerate(chunks_text)
        ]
        return ParsedDocument(
            source=str(path),
            title=path.stem,
            chunks=chunks,
            metadata={"path": str(path), "format": ext, "chunk_count": total},
        )

    def load_directory(self, directory: str | Path, glob: str = "**/*.*") -> list[ParsedDocument]:
        """Load all supported files from a directory."""
        directory = Path(directory)
        supported = {".md", ".markdown", ".txt", ".csv", ".tsv", ".pdf", ".docx"}
        docs: list[ParsedDocument] = []
        for p in directory.glob(glob):
            if p.is_file() and p.suffix.lower() in supported:
                docs.append(self.load_file(p))
        return docs

    # -- embedding + storage ---------------------------------------------------

    async def build_knowledge_base(
        self,
        source: str | Path,
        store: Any = None,
        glob: str = "**/*.*",
    ) -> int:
        """Parse, chunk, embed, and upsert all documents into a vector store.

        Returns:
            Total number of chunks inserted.
        """
        from morainet.memory.stores import InMemoryVectorStore

        if store is None:
            store = InMemoryVectorStore()
        if self.embedder is None:
            from morainet.memory.embeddings import HashEmbedder

            self.embedder = HashEmbedder()

        source_path = Path(source)
        docs = (
            self.load_directory(source_path, glob)
            if source_path.is_dir()
            else [self.load_file(source_path)]
        )

        total = 0
        for doc in docs:
            for chunk in doc.chunks:
                embedding = await self.embedder.embed(chunk.text)
                meta = {**doc.metadata, "title": doc.title, "chunk_index": chunk.chunk_index}
                await store.upsert(chunk.text, embedding, meta)
                total += 1
        return total
